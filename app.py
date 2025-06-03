from docx import Document
from deep_translator import GoogleTranslator

# Caminhos dos arquivos
input_path = "C:/Users/placido.junior/Downloads/python/Manual.docx"
output_path = "C:/Users/placido.junior/Downloads/python/Manual_Translated.docx"

# Tradutor do português para o inglês
translator = GoogleTranslator(source='pt', target='en')

# Carrega o documento
doc = Document(input_path)

# Traduz os parágrafos
for paragraph in doc.paragraphs:
    if paragraph.text.strip():
        try:
            translated = translator.translate(paragraph.text)
            paragraph.text = translated
        except Exception as e:
            print(f"Erro ao traduzir parágrafo: {paragraph.text}\n{e}")

# Traduz o conteúdo das tabelas (se houver)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if cell.text.strip():
                try:
                    cell.text = translator.translate(cell.text)
                except Exception as e:
                    print(f"Erro ao traduzir célula da tabela: {cell.text}\n{e}")

# Salva o documento traduzido
doc.save(output_path)

print("✅ Tradução completa! Arquivo salvo em:")
print(output_path)
